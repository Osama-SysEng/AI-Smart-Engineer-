"""Document processing service."""
import os
from pathlib import Path
from typing import Dict, List

from src.core.config import get_settings
from src.core.logging import get_logger
from src.db.session import async_session
from src.db.models.document import Document, DocumentPage
from src.storage.service import StorageService

logger = get_logger(__name__)


class DocumentProcessor:
    """Process uploaded documents through the pipeline."""

    def __init__(self):
        self.storage = StorageService()
        self.settings = get_settings()

    async def process(self, document_id: str) -> Dict:
        """Process document through full pipeline."""
        async with async_session() as db:
            # Get document
            from sqlalchemy import select
            result = await db.execute(select(Document).where(Document.id == document_id))
            doc = result.scalar_one_or_none()
            if not doc:
                logger.error("Document not found", document_id=document_id)
                return {"status": "error", "message": "Document not found"}

            # Update status
            doc.status = "processing"
            doc.processing_progress = 10
            await db.commit()

            try:
                # Step 1: Virus scan (simulated)
                doc.virus_scanned = True
                doc.virus_clean = True
                doc.processing_progress = 20
                await db.commit()

                # Step 2: File validation
                file_path = doc.storage_path
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"File not found: {file_path}")
                doc.processing_progress = 30
                await db.commit()

                # Step 3: OCR / Text extraction based on file type
                if doc.file_type in [".pdf", ".png", ".jpg", ".jpeg", ".tiff"]:
                    await self._process_image_document(doc, db)
                elif doc.file_type in [".xlsx", ".xls", ".csv"]:
                    await self._process_spreadsheet(doc, db)
                elif doc.file_type in [".docx", ".doc"]:
                    await self._process_word_document(doc, db)

                doc.processing_progress = 80
                await db.commit()

                # Step 4: AI extraction
                from src.services.extraction_engine import ExtractionEngine
                engine = ExtractionEngine()
                await engine.extract_from_document(doc, db)

                doc.status = "completed"
                doc.processing_progress = 100
                await db.commit()

                logger.info("Document processing completed", document_id=document_id)
                return {"status": "completed", "document_id": document_id}

            except Exception as e:
                doc.status = "failed"
                doc.error_count += 1
                await db.commit()
                logger.error("Document processing failed", error=str(e), document_id=document_id)
                return {"status": "error", "message": str(e)}

    async def _process_image_document(self, doc: Document, db) -> None:
        """Process image-based documents (PDF, images)."""
        import fitz  # PyMuPDF

        file_path = doc.storage_path

        if doc.file_type == ".pdf":
            pdf_document = fitz.open(file_path)
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]

                # Extract text
                text = page.get_text()

                # Extract images
                pix = page.get_pixmap(dpi=150)
                img_path = f"{file_path}_page_{page_num+1}.png"
                pix.save(img_path)

                # Create page record
                page_record = DocumentPage(
                    document_id=doc.id,
                    page_number=page_num + 1,
                    ocr_text=text,
                    image_path=img_path,
                    width=page.rect.width,
                    height=page.rect.height,
                )
                db.add(page_record)

            pdf_document.close()
        else:
            # Image file - use OCR
            import pytesseract
            from PIL import Image

            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang="ara+eng")

            page_record = DocumentPage(
                document_id=doc.id,
                page_number=1,
                ocr_text=text,
                image_path=file_path,
            )
            db.add(page_record)

        await db.commit()

    async def _process_spreadsheet(self, doc: Document, db) -> None:
        """Process Excel/CSV files."""
        import pandas as pd

        file_path = doc.storage_path

        if doc.file_type in [".xlsx", ".xls"]:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                # Store sheet data as metadata
                if not doc.metadata:
                    doc.metadata = {}
                doc.metadata[f"sheet_{sheet_name}"] = {
                    "columns": df.columns.tolist(),
                    "row_count": len(df),
                    "preview": df.head(5).to_dict(),
                }
        elif doc.file_type == ".csv":
            df = pd.read_csv(file_path)
            doc.metadata = {
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "preview": df.head(5).to_dict(),
            }

        await db.commit()

    async def _process_word_document(self, doc: Document, db) -> None:
        """Process Word documents."""
        from docx import Document as DocxDocument

        file_path = doc.storage_path
        word_doc = DocxDocument(file_path)

        full_text = []
        for para in word_doc.paragraphs:
            full_text.append(para.text)

        # Extract tables
        tables_data = []
        for table in word_doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)

        doc.metadata = {
            "text": "\n".join(full_text),
            "tables": tables_data,
            "paragraph_count": len(word_doc.paragraphs),
        }

        await db.commit()
